# app/utils/pdf_generator.py
from docx import Document
import os
import subprocess
import sys
import logging
from datetime import datetime
from pathlib import Path

from server.models.clan import Clan
from server.models.committee import HaiaCommittee, MadaehCommittee
from server.models.county import County
from server.models.user import User

logger = logging.getLogger(__name__)


def replace_placeholder_in_runs(runs, key, value):
    """Replace placeholder text inside runs while preserving formatting."""
    placeholder = f"{{{{{key}}}}}"
    # Merge all run text into one string
    full_text = "".join(run.text for run in runs)
    if placeholder in full_text:
        full_text = full_text.replace(placeholder, str(value or ""))
        # Clear all runs first
        for run in runs:
            run.text = ""
        # Put replaced text back into first run
        runs[0].text = full_text


def fill_docx_template(template_path: str, output_path: str, context: dict):
    """Fill DOCX template with context data."""
    try:
        doc = Document(template_path)

        # Replace in paragraphs
        for p in doc.paragraphs:
            for key, value in context.items():
                replace_placeholder_in_runs(p.runs, key, value)

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for key, value in context.items():
                            replace_placeholder_in_runs(p.runs, key, value)

        doc.save(output_path)
        logger.info(f"تم ملء قالب DOCX بنجاح: {output_path}")
    except Exception as e:
        logger.error(f"خطأ في ملء قالب DOCX: {e}")
        raise


def find_libreoffice():
    """Find LibreOffice executable path."""
    possible_paths = [
        "libreoffice",  # Linux/Mac in PATH
        "/usr/bin/libreoffice",  # Linux
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",  # Mac
        r"C:\Program Files\LibreOffice\program\soffice.exe",  # Windows
        # Windows 32-bit
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]

    for path in possible_paths:
        try:
            result = subprocess.run([path, "--version"],
                                    capture_output=True, text=True, timeout=10)
            if result.returncode == 0:
                logger.info(f"تم العثور على LibreOffice في: {path}")
                return path
        except (subprocess.TimeoutExpired, FileNotFoundError, OSError):
            continue

    return None


def convert_to_pdf(docx_path: str, pdf_path: str):
    """Convert DOCX to PDF using LibreOffice."""
    docx_path = Path(docx_path).resolve()
    pdf_path = Path(pdf_path).resolve()

    # Ensure output directory exists
    pdf_path.parent.mkdir(parents=True, exist_ok=True)

    logger.info(f"جارٍ تحويل DOCX إلى PDF...")
    logger.info(f"  المدخل: {docx_path}")
    logger.info(f"  المخرج: {pdf_path}")

    # Find LibreOffice
    libreoffice_path = find_libreoffice()
    if not libreoffice_path:
        raise Exception(
            "لم يتم العثور على LibreOffice. الرجاء تثبيت LibreOffice.")

    try:
        # Use LibreOffice to convert
        # Output will be created in the same directory as input with .pdf extension
        cmd = [
            libreoffice_path,
            "--headless",
            "--convert-to", "pdf",
            "--outdir", str(pdf_path.parent),
            str(docx_path)
        ]

        logger.info(f"تنفيذ الأمر: {' '.join(cmd)}")

        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=120  # Increased timeout to 2 minutes
        )

        # Log the output for debugging
        if result.stdout:
            logger.info(f"مخرجات LibreOffice: {result.stdout}")
        if result.stderr:
            logger.warning(f"تحذيرات LibreOffice: {result.stderr}")

        # LibreOffice creates the PDF with the same base name as the DOCX
        expected_pdf = pdf_path.parent / f"{docx_path.stem}.pdf"

        logger.info(f"موقع PDF المتوقع: {expected_pdf}")
        logger.info(f"موقع PDF المستهدف: {pdf_path}")

        # Check if conversion was successful
        if result.returncode != 0:
            error_msg = result.stderr or result.stdout or "خطأ غير معروف"
            raise Exception(f"فشل تحويل LibreOffice: {error_msg}")

        # Wait a moment for file system to sync
        import time
        time.sleep(1)

        # If the expected PDF location is different from target, move it
        if expected_pdf != pdf_path:
            if expected_pdf.exists():
                import shutil
                shutil.move(str(expected_pdf), str(pdf_path))
                logger.info(f"تم نقل PDF من {expected_pdf} إلى {pdf_path}")
            else:
                raise Exception(
                    f"لم يتم إنشاء PDF في الموقع المتوقع: {expected_pdf}")

        # Final verification
        if not pdf_path.exists():
            raise Exception(f"لم يتم إنشاء ملف PDF في: {pdf_path}")

        logger.info(f"تم التحويل إلى PDF بنجاح: {pdf_path}")

    except subprocess.TimeoutExpired:
        logger.error("انتهت مهلة تحويل LibreOffice")
        raise Exception("انتهت مهلة تحويل PDF (تجاوزت دقيقتين)")
    except Exception as e:
        logger.error(f"خطأ في تحويل LibreOffice: {e}")
        raise Exception(f"فشل تحويل PDF: {str(e)}")


def find_libreoffice():
    """Find LibreOffice executable path."""
    possible_paths = [
        "libreoffice",  # Linux in PATH
        "/usr/bin/libreoffice",  # Linux standard location
        "/usr/bin/soffice",  # Alternative Linux location
        "soffice",  # Linux in PATH
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",  # Mac
        r"C:\Program Files\LibreOffice\program\soffice.exe",  # Windows
        # Windows 32-bit
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]

    for path in possible_paths:
        try:
            result = subprocess.run(
                [path, "--version"],
                capture_output=True,
                text=True,
                timeout=10
            )
            if result.returncode == 0:
                logger.info(f"تم العثور على LibreOffice في: {path}")
                logger.info(f"الإصدار: {result.stdout.strip()}")
                return path
        except (subprocess.TimeoutExpired, FileNotFoundError, OSError):
            continue

    logger.error("لم يتم العثور على LibreOffice في أي موقع قياسي")
    return None


def find_template_path():

    # Get the directory where pdf_generator.py is located (server/utils/)
    current_file_dir = Path(__file__).parent  # server/utils/

    # Go up one level to server/, then navigate to app/templates/
    server_dir = current_file_dir.parent  # server/
    template_path = server_dir / "app" / "templates" / "wedding_request_form.docx"

    if template_path.exists():
        logger.info(f"تم العثور على القالب في: {template_path.resolve()}")
        return str(template_path.resolve())

    # Fallback: Log the issue for debugging
    logger.error(
        f"لم يتم العثور على القالب في المسار المتوقع: {template_path.resolve()}")
    logger.error(f"موقع الملف الحالي: {Path(__file__).resolve()}")
    logger.error(f"مجلد الخادم: {server_dir.resolve()}")

    raise FileNotFoundError(
        f"لم يتم العثور على wedding_request_form.docx في {template_path.resolve()}"
    )


def generate_wedding_pdf(reservation, output_dir: str, db):
    """Generate wedding PDF from reservation data."""
    try:
        # Ensure output directory exists
        output_dir = Path(output_dir).resolve()
        output_dir.mkdir(parents=True, exist_ok=True)

        # Find template
        template_path = find_template_path()

        # Generate unique file paths with timestamp to avoid conflicts
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        unique_id = f"{reservation.id}_{timestamp}"

        filled_docx_path = output_dir / f"reservation_{unique_id}_filled.docx"
        pdf_path = output_dir / f"reservation_{unique_id}.pdf"

        logger.info(f"جارٍ إنشاء PDF للحجز {reservation.id}")
        logger.info(f"القالب: {template_path}")
        logger.info(f"مخرج DOCX: {filled_docx_path}")
        logger.info(f"مخرج PDF: {pdf_path}")

        # Get database values with error handling
        try:
            user_of_this_reservation = db.query(User).filter(
                User.id == reservation.groom_id).first()
            if not user_of_this_reservation:
                raise Exception(
                    f"لم يتم العثور على المستخدم للعريس: {reservation.groom_id}")

            reserved_clan = db.query(Clan).filter(
                Clan.id == reservation.clan_id).first()
            if not reserved_clan:
                raise Exception(
                    f"لم يتم العثور على العشيرة المحجوزة: {reservation.clan_id}")

            original_clan = db.query(Clan).filter(
                Clan.id == user_of_this_reservation.clan_id).first()
            if not original_clan:
                raise Exception(
                    f"لم يتم العثور على العشيرة الأصلية: {user_of_this_reservation.clan_id}")

            county = db.query(County).filter(
                County.id == user_of_this_reservation.county_id).first()
            if not county:
                raise Exception(
                    f"لم يتم العثور على البلدية: {user_of_this_reservation.county_id}")

            haia_committee = db.query(HaiaCommittee).filter(
                HaiaCommittee.id == reservation.haia_committee_id).first()

            madaeh_committee = db.query(MadaehCommittee).filter(
                MadaehCommittee.id == reservation.madaeh_committee_id).first()

        except Exception as e:
            logger.error(f"خطأ في استعلام قاعدة البيانات: {e}")
            raise Exception(f"فشل في جلب البيانات المطلوبة: {e}")

        # Prepare context data
        context = {
            "COUNTY": county.name if county else "",
            "ORIGIN_CLAN": original_clan.name if original_clan else "",
            "RESERVED_CLAN": reserved_clan.name if reserved_clan else "",
            "groom_NAME": user_of_this_reservation.first_name or "",
            "last_name": user_of_this_reservation.last_name or "",
            "GUARDIAN_NAME": user_of_this_reservation.guardian_name or "",
            "father_name": user_of_this_reservation.father_name or "",
            "guardian_birth_date": user_of_this_reservation.guardian_birth_date.strftime("%Y-%m-%d") if user_of_this_reservation.guardian_birth_date else "",
            "guardian_birth_address": user_of_this_reservation.guardian_birth_address or "",
            "guardian_home_address": user_of_this_reservation.guardian_home_address or "",
            "grandfather_name": user_of_this_reservation.grandfather_name or "",
            "birth_date": user_of_this_reservation.birth_date.strftime("%Y-%m-%d") if user_of_this_reservation.birth_date else "",
            "birth_address": user_of_this_reservation.birth_address or "",
            "home_address": user_of_this_reservation.home_address or "",
            "phone_number": user_of_this_reservation.phone_number or "",
            "WEDDING_DATES": f"{reservation.date1.strftime('%Y-%m-%d')} - {reservation.date2.strftime('%Y-%m-%d')}" if reservation.date2 else reservation.date1.strftime("%Y-%m-%d"),
            "haia_committee_id": haia_committee.name if haia_committee else "",
            "madaeh_committee_id": madaeh_committee.name if madaeh_committee else "",
            "GUARDIAN_phone": user_of_this_reservation.guardian_phone or "",
            "created_at": reservation.created_at.strftime("%Y-%m-%d") if reservation.created_at else "",
        }

        # Fill DOCX template
        fill_docx_template(template_path, str(filled_docx_path), context)

        # Verify DOCX was created
        if not filled_docx_path.exists():
            raise Exception("لم يتم إنشاء ملف DOCX بنجاح")

        logger.info(f"تم إنشاء DOCX بنجاح: {filled_docx_path}")

        # Convert to PDF
        convert_to_pdf(str(filled_docx_path), str(pdf_path))

        # Verify PDF was created
        if not pdf_path.exists():
            raise Exception("لم يتم إنشاء ملف PDF بنجاح")

        logger.info(f"تم إنشاء PDF بنجاح: {pdf_path}")

        # Always clean up the intermediate DOCX file
        try:
            if filled_docx_path.exists():
                filled_docx_path.unlink()
                logger.info(
                    f"تم حذف ملف DOCX المؤقت: {filled_docx_path}")
        except Exception as e:
            logger.warning(f"لم يتمكن من حذف DOCX المؤقت: {e}")

        return str(pdf_path)

    except Exception as e:
        logger.error(
            f"فشل إنشاء PDF للحجز {reservation.id}: {e}")
        # Clean up any partial files
        try:
            if 'filled_docx_path' in locals() and filled_docx_path.exists():
                filled_docx_path.unlink()
                logger.info("تم حذف ملف DOCX الجزئي")
            if 'pdf_path' in locals() and pdf_path.exists():
                pdf_path.unlink()
                logger.info("تم حذف ملف PDF الجزئي")
        except Exception as cleanup_error:
            logger.warning(f"خطأ في التنظيف: {cleanup_error}")

        raise Exception(f"فشل إنشاء PDF: {e}")


# Additional utility function for testing
def test_pdf_generation():
    """Test PDF generation capabilities."""
    print("اختبار قدرات إنشاء PDF...")

    # Test template
    try:
        template_path = find_template_path()
        print(f"✓ تم العثور على القالب: {template_path}")
    except Exception as e:
        print(f"✗ لم يتم العثور على القالب: {e}")
        return False

    # Test LibreOffice
    libreoffice_path = find_libreoffice()
    if libreoffice_path:
        print(f"✓ تم العثور على LibreOffice: {libreoffice_path}")
    else:
        print("✗ لم يتم العثور على LibreOffice")

    # Test docx2pdf
    try:
        import docx2pdf
        print("✓ docx2pdf متاح")
    except ImportError:
        print("✗ docx2pdf غير متاح")

    return True


if __name__ == "__main__":
    test_pdf_generation()
